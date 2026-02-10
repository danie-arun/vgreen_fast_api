from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import logging

logger = logging.getLogger(__name__)


class ExportService:
    @staticmethod
    def _set_cell_background(cell, fill):
        """Set cell background color"""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), fill)
        cell._element.get_or_add_tcPr().append(shading_elm)

    @staticmethod
    def _add_table_to_document(doc, title, columns, rows):
        """Add a table to the document with proper formatting"""
        doc.add_heading(title, level=2)
        
        table = doc.add_table(rows=1, cols=len(columns))
        table.style = 'Light Grid Accent 1'
        
        # Add header row
        header_cells = table.rows[0].cells
        for i, col in enumerate(columns):
            header_cells[i].text = col
            # Style header
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ExportService._set_cell_background(header_cells[i], '0070C0')
        
        # Add data rows
        for row_data in rows:
            row_cells = table.add_row().cells
            for i, value in enumerate(row_data):
                row_cells[i].text = str(value)
                # Align right for numeric columns
                for paragraph in row_cells[i].paragraphs:
                    if i > 0 and isinstance(value, (int, float)):
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        doc.add_paragraph()  # Add spacing

    @staticmethod
    def export_financial_summary(summary_data, metrics):
        """Export Financial Summary table to Word document"""
        try:
            doc = Document()
            
            # Add title and metadata
            title = doc.add_heading('Financial Summary Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            doc.add_paragraph()
            
            # Add metrics summary
            doc.add_heading('Key Metrics', level=1)
            metrics_table = doc.add_table(rows=1, cols=2)
            metrics_table.style = 'Light Grid Accent 1'
            
            header_cells = metrics_table.rows[0].cells
            header_cells[0].text = 'Metric'
            header_cells[1].text = 'Value'
            
            for key, value in metrics.items():
                row_cells = metrics_table.add_row().cells
                row_cells[0].text = key
                row_cells[1].text = str(value)
            
            doc.add_paragraph()
            
            # Add financial summary table
            columns = ['Loan ID', 'Group Name', 'Loan Amount (₹)', 'Collected (₹)', 'Pending (₹)', 'EMI Day', 'Status']
            rows = []
            
            for item in summary_data:
                # loanAmount is now a string like "principal + interest", use loanAmountValue for numeric
                loan_amt = item.get('loanAmountValue', 0)
                if isinstance(loan_amt, str):
                    loan_amt = 0
                
                rows.append([
                    item.get('loanId', ''),
                    item.get('groupName', ''),
                    f"₹{float(loan_amt):,.2f}",
                    f"₹{item.get('collectedAmount', 0):,.2f}",
                    f"₹{item.get('pendingAmount', 0):,.2f}",
                    item.get('emiDay', ''),
                    item.get('status', ''),
                ])
            
            ExportService._add_table_to_document(doc, 'Financial Summary', columns, rows)
            
            return doc
        except Exception as e:
            logger.exception(f"Error exporting financial summary: {str(e)}")
            raise

    @staticmethod
    def export_user_summary(user_summary_data):
        """Export User Summary table to Word document"""
        try:
            doc = Document()
            
            # Add title and metadata
            title = doc.add_heading('User Summary Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            doc.add_paragraph()
            
            # Add user summary table
            columns = ['User Name', 'Loan ID', 'Group Name', 'Total EMI (₹)', 'Paid EMI (₹)', 'Pending EMI (₹)']
            rows = []
            
            for item in user_summary_data:
                rows.append([
                    item.get('userName', ''),
                    item.get('loanId', ''),
                    item.get('groupName', ''),
                    f"₹{item.get('totalEmi', 0):,.2f}",
                    f"₹{item.get('paidEmi', 0):,.2f}",
                    f"₹{item.get('pendingEmi', 0):,.2f}",
                ])
            
            ExportService._add_table_to_document(doc, 'User Summary', columns, rows)
            
            return doc
        except Exception as e:
            logger.exception(f"Error exporting user summary: {str(e)}")
            raise

    @staticmethod
    def export_emi_summary(emi_summary_data):
        """Export EMI Summary table with details to Word document"""
        try:
            doc = Document()
            
            # Add title and metadata
            title = doc.add_heading('EMI Summary Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            doc.add_paragraph()
            
            # Add EMI summary table
            columns = ['Loan ID', 'User Name', 'Total EMIs', 'Paid EMIs', 'Pending EMIs']
            rows = []
            
            for item in emi_summary_data:
                rows.append([
                    item.get('loanId', ''),
                    item.get('userName', ''),
                    item.get('totalEmis', 0),
                    item.get('paidEmis', 0),
                    item.get('pendingEmis', 0),
                ])
            
            ExportService._add_table_to_document(doc, 'EMI Summary', columns, rows)
            
            # Add detailed EMI information
            doc.add_heading('EMI Details', level=1)
            
            for item in emi_summary_data:
                doc.add_heading(f"Loan: {item.get('loanId')} - User: {item.get('userName')}", level=2)
                
                emi_details = item.get('emiDetails', [])
                if emi_details:
                    detail_columns = ['EMI Date', 'EMI Amount (₹)', 'Status']
                    detail_rows = []
                    
                    for emi in emi_details:
                        detail_rows.append([
                            emi.get('emiDate', ''),
                            f"₹{emi.get('emiAmount', 0):,.2f}",
                            emi.get('emiStatus', ''),
                        ])
                    
                    ExportService._add_table_to_document(doc, f"EMI Details for {item.get('loanId')}", detail_columns, detail_rows)
            
            return doc
        except Exception as e:
            logger.exception(f"Error exporting EMI summary: {str(e)}")
            raise

    @staticmethod
    def export_collections_summary(collections_summary_data):
        """Export Collections Summary table with detailed user information to Word document"""
        try:
            doc = Document()
            
            # Set document to landscape orientation
            section = doc.sections[0]
            section.page_height = Inches(8.5)
            section.page_width = Inches(11)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
         
            # Add collections summary table
            columns = ['Loan Number', 'Group Name', 'Loan Amount (₹)', 'Collected Amount (₹)', 'Pending Amount (₹)', 'Next EMI Date', 'Next EMI Amount (₹)']
            rows = []
            
            for item in collections_summary_data:
                # loanAmount is now a string like "principal + interest", use loanAmountValue for numeric
                loan_amt = item.get('loanAmountValue', 0)
                if isinstance(loan_amt, str):
                    loan_amt = 0
                
                rows.append([
                    item.get('loanId', ''),
                    item.get('groupName', ''),
                    f"₹{float(loan_amt):,.2f}",
                    f"₹{item.get('collectedAmount', 0):,.2f}",
                    f"₹{item.get('pendingAmount', 0):,.2f}",
                    item.get('nextEmiDate', ''),
                    f"₹{item.get('nextEmiAmount', 0):,.2f}",
                ])
            
            #ExportService._add_table_to_document(doc, 'Collections Summary', columns, rows)
            
            first_loan = True
            for item in collections_summary_data:
                user_details = item.get('userDetails', [])
                if user_details:
                    # Add page break before each loan (except the first)
                    if not first_loan:
                        doc.add_page_break()
                    first_loan = False
                    
                    # Add title and generated date to each page
                    title = doc.add_heading('Collections Summary Report', 0)
                    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
                    doc.add_paragraph()
                    
                    # Add loan header info table
                    header_table = doc.add_table(rows=2, cols=2)
                    header_table.style = 'Light Grid Accent 1'
                    
                    # Row 1: Loan ID and Group Name
                    header_table.rows[0].cells[0].text = f"Loan ID: {item.get('loanId', 'N/A')}"
                    header_table.rows[0].cells[1].text = f"Group Name: {item.get('groupName', 'N/A')}"
                    
                    # Row 2: EMI Day and Collection Date
                    header_table.rows[1].cells[0].text = f"EMI Day: {item.get('emiDay', 'N/A')}"
                    header_table.rows[1].cells[1].text = f"Collection Date: {item.get('collectionDate', 'N/A')}"
                    
                    # Format header table
                    for row in header_table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(10)
                                    run.font.bold = True
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    doc.add_paragraph()  # Add spacing
                    
                    # Create table with users and their details
                    # Columns: Member, Mobile, LA, Paid, Pending, EMI Adv, No EMI, No Paid, No Overdue, OD amt, Loan Adv, EMI amt
                    columns = ['Member', 'Mobile', 'LA', 'Paid', 'Pending', 'EMI Adv', 'No EMI', 'No Paid', 'No OD', 'OD Amt', 'Loan Adv', 'EMI Amt']
                    
                    user_table = doc.add_table(rows=1, cols=len(columns))
                    user_table.style = 'Light Grid Accent 1'
                    user_table.autofit = False
                    user_table.allow_autofit = False
                    
                    # Set table width to full page
                    table_width = Inches(10)
                    user_table.width = table_width
                    
                    # Add header row
                    header_cells = user_table.rows[0].cells
                    for i, col in enumerate(columns):
                        header_cells[i].text = col
                        # Style header with smaller font
                        for paragraph in header_cells[i].paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                                run.font.size = Pt(9)
                                run.font.color.rgb = RGBColor(255, 255, 255)
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        ExportService._set_cell_background(header_cells[i], '0070C0')
                    
                    # Calculate totals for this loan
                    total_no_od = 0
                    total_od_amt = 0
                    total_loan_adv = 0
                    total_emi_amt = 0
                    
                    # Add user rows
                    for user in user_details:
                        row_cells = user_table.add_row().cells
                        
                        # Set cell values (without ₹ symbols)
                        row_cells[0].text = user.get('userName', 'N/A')
                        row_cells[1].text = user.get('mobileNumber', 'N/A')
                        row_cells[2].text = f"{user.get('loanAmount', 0):,.0f}"
                        row_cells[3].text = f"{user.get('collectedAmount', 0):,.0f}"
                        row_cells[4].text = f"{user.get('pendingAmount', 0):,.0f}"
                        row_cells[5].text = f"{user.get('loanAdvance', 0):,.0f}"  # EMI Adv = loan member advance
                        row_cells[6].text = str(user.get('totalEmis', 0))
                        row_cells[7].text = str(user.get('paidEmis', 0))
                        row_cells[8].text = str(user.get('overdueEmis', 0))  # No OD = count of EMIs with date < today
                        row_cells[9].text = f"{user.get('totalOverdueAmount', 0):,.0f}"  # OD Amt = sum of EMI amounts with date < today
                        row_cells[10].text = "100"  # Loan Adv = loan member advance
                        row_cells[11].text = f"{user.get('emiAmount', 0):,.0f}"
                        
                        # Accumulate totals
                        total_no_od += user.get('overdueEmis', 0)
                        total_od_amt += user.get('totalOverdueAmount', 0)
                        total_loan_adv += 100  # Loan Adv is fixed at 100
                        total_emi_amt += user.get('emiAmount', 0)
                        
                        # Format all cells with smaller font and proper alignment
                        for i, cell in enumerate(row_cells):
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(8)
                                # Align numeric columns to right, text columns to left
                                if i == 0 or i == 1:
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                else:
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Add total row
                    total_row_cells = user_table.add_row().cells
                    total_row_cells[0].text = "Total Amount"
                    
                    # Make "Total Amount" bold and span across Member to No Paid columns (0-7)
                    for i in range(8):
                        if i > 0:
                            total_row_cells[i].text = ""
                    
                    # Format total row header
                    for paragraph in total_row_cells[0].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(8)
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    # Add totals for No OD, OD Amt, Loan Adv, EMI Amt
                    total_row_cells[8].text = str(total_no_od)
                    total_row_cells[9].text = f"{total_od_amt:,.0f}"
                    total_row_cells[10].text = f"{total_loan_adv:,.0f}"
                    total_row_cells[11].text = f"{total_emi_amt:,.0f}"
                    
                    # Format total row cells
                    for i in range(8, 12):
                        for paragraph in total_row_cells[i].paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                                run.font.size = Pt(8)
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    doc.add_paragraph()  # Add spacing
                    
                    # Add summary values below table
                    summary_para = doc.add_paragraph()
                    summary_para.add_run(f"Total EMI Amount Collection: {total_emi_amt:,.0f}\n")
                    summary_para.add_run(f"Total Advance Collection: {total_loan_adv:,.0f}\n")
                    summary_para.add_run(f"Total Over Due Collection: {total_od_amt:,.0f}")
                    
                    for run in summary_para.runs:
                        run.font.size = Pt(9)
                    
                    doc.add_paragraph()  # Add spacing
                    
                    # Add signature areas
                    sig_table = doc.add_table(rows=2, cols=2)
                    sig_table.style = 'Light Grid Accent 1'
                    
                    # Header row
                    sig_table.rows[0].cells[0].text = "Employer Signature"
                    sig_table.rows[0].cells[1].text = "Manager Signature"
                    
                    # Signature lines (empty rows for signing)
                    sig_table.rows[1].cells[0].text = "_" * 30
                    sig_table.rows[1].cells[1].text = "_" * 30
                    
                    # Format signature table
                    for row in sig_table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(9)
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            return doc
        except Exception as e:
            logger.exception(f"Error exporting collections summary: {str(e)}")
            raise
